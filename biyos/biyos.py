import cookielib
import urllib
import urllib2
import sys
import biyosui
import base64 as b64
import re

from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook

from PyQt4 import QtGui
from bs4 import BeautifulSoup
from math import ceil

# No, Blok, Daire
kiraci = [[7710, "A", 6]]


class BiyosApp(QtGui.QMainWindow, biyosui.Ui_MainWindow):
    def __init__(self, parent=None):
        super(BiyosApp, self).__init__(parent)
        self.setupUi(self)
        self.dogalgaz_birim_in.setValue(11)
        self.su_birim_in.setValue(5)

        self.kalori_hesap_button.clicked.connect(self.kalori_hesap)
        self.sayac_veri_button.clicked.connect(self.sayac_verileri)
        self.apartman_aidat_button.clicked.connect(self.apartman_aidat)
        self.tum_borclar_button.clicked.connect(self.tum_borclar)
        self.tek_borc_button.clicked.connect(self.tek_borc)
        self.login()

    def login(self):
        with open('login/log.in', 'r') as f:
            self.email = b64.decodestring(f.readline().strip())
            self.password = b64.decodestring(f.readline().strip())

        self.cj = cookielib.CookieJar()
        self.opener = urllib2.build_opener(
            urllib2.HTTPRedirectHandler(),
            urllib2.HTTPHandler(debuglevel=0),
            urllib2.HTTPSHandler(debuglevel=0),
            urllib2.HTTPCookieProcessor(self.cj)
        )

        self.opener.addheaders = [('User-agent', ('Mozilla/4.0 (compatible; MSIE 6.0; '
                                                  'Windows NT 5.2; .NET CLR 1.1.4322)'))]

        # need this twice - once to set cookies, once to log in...
        self._login()
        self._login()
        self.giris_button.setStyleSheet('QPushButton {background-color: #00FF00; color: black;}')
        self.giris_button.setText(self.email + ' adresi ile giris yapildi!')

    def _login(self):
        """
        Handle login. This should populate our cookie jar.
        """
        login_data = urllib.urlencode({
            'email': self.email,
            'password': self.password,
        })

        response = self.opener.open("https://app.biyos.net/login.php", login_data)

    def sayac_verileri(self):
        self.dogalgaz_birim = float(self.dogalgaz_birim_in.value())
        self.su_birim = float(self.su_birim_in.value())
        su = self.get_page('https://app.biyos.net/yonetim?sayac_tipi=sicaksu')
        self.su_toplam = self.get_sayac_toplam(su)
        self.su_toplam_disp.setText(str(self.su_toplam))

        kalori = self.get_page('https://app.biyos.net/yonetim?sayac_tipi=kalorimetre')
        self.kalori_toplam = self.get_sayac_toplam(kalori)
        self.kalori_toplam_disp.setText(str(self.kalori_toplam))

        self.kalori_ortalama = self.kalori_toplam / 48.0
        self.kalori_ortalama_disp.setText(str("%.2f" % self.kalori_ortalama))

        self.sayac_veri_button.setStyleSheet('QPushButton {background-color: #00FF00; color: black;}')
        self.sayac_veri_button.setText('Veriler gosteriliyor')

    def kalori_hesap(self):
        self.sayac_verileri()
        self.dogalgaz_birim = float(self.dogalgaz_birim_in.value())
        self.su_birim = float(self.su_birim_in.value())
        fatura = float(self.fatura_in.value())

        if fatura == 0:
            self.kalori_hesap_button.setStyleSheet('QPushButton {background-color: #FF0000; color: black;}')
            self.kalori_hesap_button.setText('Fatura girip tekrar deneyin!')
            return

        su_fark = (self.dogalgaz_birim - self.su_birim) * self.su_toplam
        son_fiyat = fatura - su_fark
        self.son_fiyat_disp.setText(str("%.2f" % son_fiyat))
        ortak_gider = (son_fiyat * 3.) / 480.
        aidat = 200. - ortak_gider
        self.ortak_gider_disp.setText(str("%.2f" % ortak_gider))
        self.aidat_disp.setText(str("%.2f" % aidat))

        self.kalori_hesap_button.setStyleSheet('QPushButton {background-color: #00FF00; color: black;}')
        self.kalori_hesap_button.setText('Hesaplandi!')

    def _get_rows(self, html, attr=None):
        if attr is None:
            attr = "table"

        table = html.find('table', attrs={'class': attr})
        body = table.find('tbody')
        rows = body.find_all('tr')

        return rows

    def get_sayac_toplam(self, html):
        rows = self._get_rows(html)

        total = 0
        for r in rows:
            cols = r.find_all('td')
            total += int(cols[-1].text)

        return total

    def get_page(self, url):
        try:
            resp = self.opener.open(url)
            return BeautifulSoup(resp.read(), "lxml")
        except Exception as e:
            raise e
            return

    def apartman_aidat(self):
        self.sayac_verileri()
        url = 'https://app.biyos.net/raporlar/paylasimlar/' + str(self.paylasim_link_in.value())
        su_rows = []
        kalori_rows = []
        title = ""
        try:
            kalori = self.get_page(url)
            su = self.get_page('https://app.biyos.net/yonetim?sayac_tipi=sicaksu')
            section = kalori.body.find('section', attrs={'class': 'rapor'})
            title = section.find('h4', attrs={'class': 'pull-left'}).get_text()
            yil = title.split('-')[0].strip()
            ay = title.split('-')[1].strip().split(' ')[0].strip()
            title = yil + ' - ' + ay
            su_rows = self._get_rows(su)
            kalori_rows = self._get_rows(kalori)
        except Exception as e:
            print e
            self.apartman_aidat_button.setStyleSheet('QPushButton {background-color: #FF0000; color: white;}')
            self.apartman_aidat_button.setText('Yazdirma basarisiz, linki kontrol edin!')
            return

        try:
            self.wb = load_workbook('aidat/template/aidat.xlsx')
            ws = self.wb.active
            ws.title = title
            ws['C1'] = ws['C29'] = title
            self._set_xlsx(ws, su_rows, kalori_rows)

            self.wb.save(filename='aidat/' + title + ' ISIMLI Aidat.xlsx')
            self._remove_names(ws)
            self.wb.save(filename='aidat/' + title + ' ISIMSIZ Aidat.xlsx')

        except Exception as e:
            print e
            self.apartman_aidat_button.setStyleSheet('QPushButton {background-color: #FF0000; color: white;}')
            self.apartman_aidat_button.setText('Yazdirma basarisiz!')
        else:
            self.apartman_aidat_button.setStyleSheet('QPushButton {background-color: #00FF00; color: black;}')
            self.apartman_aidat_button.setText(title + ' Yazdirildi!')

    def _remove_names(self, ws):
        for i in range(4, 28):
            ws.cell(row=i, column=2).value = 'NO LU Daire'
            ws.cell(row=i+28, column=2).value = 'NO LU Daire'

    def _set_xlsx(self, ws, su, kalori):
        for i in range(48):
            r = i + 4
            if i >= 24:
                r += 4

            col = su[i].find_all('td')
            ws.cell(row=r, column=2).value = col[2].text
            ws.cell(row=r, column=3).value = int(col[5].text)
            ws.cell(row=r, column=4).value = su_tl = self.dogalgaz_birim * int(col[5].text)

            col = kalori[i].find_all('td')
            ws.cell(row=r, column=5).value = float(col[6].text.replace(',', '.'))
            ws.cell(row=r, column=6).value = d70 = float(col[8].text.replace(',', '.'))
            ws.cell(row=r, column=7).value = d30 = float(col[7].text.replace(',', '.'))

            aidat = 200. - d30
            ws.cell(row=r, column=8).value = aidat
            total = su_tl + d70 + d30 + aidat
            ws.cell(row=r, column=9).value = ceil(total)

    def _single_account(self, no, blok, daire):
        html = self.get_page('https://app.biyos.net/hesaplar/' + str(no))
        hesap =  html.body.find('span', attrs={'style': 'font-size:22px;'}).get_text()

        head = self.document.add_heading(hesap, level=1)
        head.style.paragraph_format.keep_together = True
        head.style.paragraph_format.keep_with_next = True
        head = self.document.add_heading(blok + " Blok / No: " + str(daire), level=2)
        head.style.paragraph_format.keep_together = True
        head.style.paragraph_format.keep_with_next = True

        try:
            data = html.body.find('div', attrs={'class': 'table-responsive'})
            geciken = html.body.find('div', attrs={'class': 'detail-payment-item text-danger big-title'})
            bakiye = html.body.find('div', attrs={'class': 'detail-payment-item text-warning big-title'})
            self.create_table(data, geciken, bakiye)
        except AttributeError:
            return

    def create_table(self, data, geciken, bakiye):
        if bakiye:
            rows = self._get_rows(data, attr='table table-detail')

            tbl = self.document.add_table(rows=0, cols=3)
            tbl.autofit = False
            tbl.style.paragraph_format.keep_together = True
            tbl.style.paragraph_format.keep_with_next = True
            tbl.style.paragraph_format.widow_control = True

            row_cells = tbl.add_row().cells
            row_cells[0].text = "Son Odeme Tarihi"
            row_cells[1].text = "Aciklama"
            row_cells[2].text = "Tutar"

            for r in rows:
                row_cells = tbl.add_row().cells
                cols = r.find_all('td')
                i = 0
                for c in cols:
                    if c.text:
                        row_cells[i].text = c.text
                        i += 1

            non_decimal = re.compile(r'[^\d.,]+')

            row_cells = tbl.add_row().cells
            row_cells[1].text = "Toplam Borc"
            row_cells[2].text = non_decimal.sub('', bakiye.get_text())

            tbl.columns[0].width = Inches(1.5)
            tbl.columns[1].width = Inches(50)
            tbl.columns[2].width = Inches(0.5)

        else:
            self.document.add_heading("Odenmemis borcunuz bulunmamaktadir.", level=3)
            self.document.add_heading("Gosterdiginiz hassasiyet icin tesekkur ederiz.", level=4)

    def tek_borc(self):
        blok = None
        d = 0
        if self.a_blok_in.isChecked():
            d = 0
            blok = "A"
        elif self.b_blok_in.isChecked():
            d = 24
            blok = "B"
        else:
            self.tek_borc_button.setStyleSheet('QPushButton {background-color: #FF0000; color: white;}')
            self.tek_borc_button.setText('Blok seciniz!')
            return

        daire = int(self.daire_no_in.value())
        hesap = daire + 6148 + d
        yazdir = [[hesap, blok, daire]]
        for k in kiraci:
            if k[1] == blok and k[2] == daire:
                yazdir.append(k)

        try:
            self.document = Document()
            for d in yazdir:
                self._single_account(*d)
                self.document.save('aidat/' + d[1] + '-' + str(d[2]) + ' borc.docx')
        except Exception as e:
            print e
            self.tek_borc_button.setStyleSheet('QPushButton {background-color: #FF0000; color: white;}')
            self.tek_borc_button.setText('Yazdirma basarisiz!')
        else:
            self.tek_borc_button.setStyleSheet('QPushButton {background-color: #00FF00; color: black;}')
            self.tek_borc_button.setText('Basarili!\nBaska Yazdir')

    def tum_borclar(self):
        self.tum_borclar_button.setText('Yazdiriliyor, lutfen bekleyin...')

        try:
            self.document = Document()

            bar = "".join(['_'] * 78)

            daire = 1
            blok = "A"
            for i in range(6149, 6197):
                print blok, daire

                p = self.document.add_paragraph()
                p.add_run(bar).bold = True
                p.style.paragraph_format.keep_together = True
                p.style.paragraph_format.keep_with_next = True

                self._single_account(i, blok, daire)

                daire += 1
                if daire == 25:
                    daire = 1
                    blok = "B"

            for k in kiraci:
                p = self.document.add_paragraph()
                p.style.paragraph_format.keep_together = True
                p.style.paragraph_format.keep_with_next = True
                p.add_run(bar).bold = True

                self._single_account(*k)

            self.document.save('aidat/Tum borclar.docx')

        except Exception as e:
            print e
            self.tum_borclar_button.setStyleSheet('QPushButton {background-color: #FF0000; color: white;}')
            self.tum_borclar_button.setText('Yazdirma basarisiz!')
        else:
            self.tum_borclar_button.setStyleSheet('QPushButton {background-color: #00FF00; color: black;}')
            self.tum_borclar_button.setText('Yazdirma basarili!')


def main():
    app = QtGui.QApplication(sys.argv)
    biyos = BiyosApp()
    biyos.show()
    app.exec_()

if __name__ == '__main__':
    main()
